from pypdf import PdfReader
import docx


def clean_text(text):
    # Replace non-XML compatible characters with an empty string
    cleaned_text = ''.join(c for c in text if c.isprintable())
    return cleaned_text


def pdf_to_docx(file, name):
    reader = PdfReader(file)
    text = ""
    document = docx.Document()
    for i in range(len(reader.pages)):
        text += clean_text(reader.pages[i].extract_text())
        document.add_paragraph(text)

    document.save(f"/storage/emulated/0/Download/{name}.docx")


def create_docx_from_text(text, name):
    document = docx.Document()
    document.add_paragraph(text)

    document.save(f"storage/emulated/0/Download/Pro Scanner/docx/{name}.docx")
