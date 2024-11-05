from docx import Document
from docx.shared import RGBColor

def export_docx(content, name,path):
    # Create a new Document
    doc = Document()

    # Add a title with black color
#   title_paragraph = doc.add_heading(title, level=1)
   # run = title_paragraph.runs[0]  # Get the run of the title
    #run.font.color.rgb = RGBColor(0, 0, 0)  # Set title color to black

    # Split content by new lines and add to the document
    for line in content.splitlines():
           content_paragraph = doc.add_paragraph(line)
           for run in content_paragraph.runs:
               run.font.color.rgb = RGBColor(0, 0, 0) # Set content color to black

    # Save the document
    try:
        doc.save(path)
        return "Success"
    except:
        return "Failure"

